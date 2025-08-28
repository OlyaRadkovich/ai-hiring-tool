import os
import json
import io
import re
import asyncio
from typing import Tuple
from loguru import logger

from backend.api.models import PreparationAnalysis, ResultsAnalysis, ScoreBreakdown
from ..core.config import settings

from backend.agents.pipeline_1_pre_interview.agent_1_data_parser import agent_1_data_parser
from backend.agents.pipeline_1_pre_interview.agent_2_grader import agent_2_grader
from backend.agents.pipeline_1_pre_interview.agent_3_report_generator import agent_3_report_generator
from backend.agents.pipeline_2_post_interview.agent_4_topic_extractor import agent_4_topic_extractor
from backend.agents.pipeline_2_post_interview.agent_5_final_report_generator import agent_5_final_report_generator

import assemblyai as aai
from pypdf import PdfReader
import docx
from google.adk.runners import Runner
from google.adk.sessions import InMemorySessionService
from google.genai import types
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from docx import Document
from docx.shared import Pt


class AnalysisService:
    """Service responsible for interview analysis business logic using AI Agents"""

    def __init__(self):
        if settings.assemblyai_api_key:
            aai.settings.api_key = settings.assemblyai_api_key
            logger.success("Клиент AssemblyAI сконфигурирован.")
        else:
            logger.warning("API ключ для AssemblyAI не настроен в .env файле!")

        self.drive_service = None
        try:
            credentials_path = settings.google_application_credentials
            if os.path.exists(credentials_path):
                credentials = service_account.Credentials.from_service_account_file(credentials_path)
                self.drive_service = build('drive', 'v3', credentials=credentials)
                logger.success("Клиент Google Drive API успешно инициализирован.")
            else:
                logger.error(f"Файл учетных данных Google не найден по пути: {credentials_path}")
        except Exception as e:
            logger.error(f"Ошибка инициализации клиента Google Drive API: {e}")

    def _set_google_api_key(self):
        """
        Устанавливает ключ Google API как переменную окружения.
        """
        api_key_to_use = settings.google_api_key
        if not api_key_to_use:
            logger.error("Ключ Google API (google_api_key) не найден в .env файле.")
            raise ValueError("Google API key is not provided.")
        os.environ['GOOGLE_API_KEY'] = api_key_to_use
        logger.info("Ключ Google API установлен как переменная окружения для текущего запроса.")

    def _get_google_drive_file_id(self, link: str) -> str:
        """
        Извлекает ID файла из ссылки на Google Drive.
        """
        patterns = [
            r'/file/d/([a-zA-Z0-9_-]+)',
            r'/spreadsheets/d/([a-zA-Z0-9_-]+)'
        ]
        for pattern in patterns:
            match = re.search(pattern, link)
            if match:
                return match.group(1)
        raise ValueError("Некорректная ссылка на Google Drive. Не удалось извлечь ID файла.")

    async def _download_sheet_from_drive(self, file_id: str) -> str:
        """
        Загружает Google Таблицу как CSV и возвращает ее текстовое содержимое.
        """
        if not self.drive_service:
            raise ConnectionError("Сервис Google Drive не инициализирован.")
        logger.info(f"Начало загрузки таблицы с ID: {file_id} из Google Drive.")
        try:
            request = self.drive_service.files().export_media(fileId=file_id, mimeType='text/csv')
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)

            def download_in_thread():
                done = False
                while not done:
                    _, done = downloader.next_chunk()

            await asyncio.to_thread(download_in_thread)
            logger.success(f"Таблица {file_id} успешно экспортирована в CSV.")
            file_io.seek(0)
            return file_io.read().decode('utf-8')
        except Exception as e:
            logger.error(f"Ошибка при экспорте таблицы из Google Drive: {e}", exc_info=True)
            raise IOError(f"Не удалось загрузить требования из Google Drive: {e}")

    async def _download_audio_from_drive(self, file_id: str) -> io.BytesIO:
        """
        Асинхронно загружает аудио/видео файл из Google Drive.
        """
        if not self.drive_service:
            raise ConnectionError("Сервис Google Drive не инициализирован. Проверьте учетные данные.")
        logger.info(f"Начало загрузки файла с ID: {file_id} из Google Drive.")
        try:
            request = self.drive_service.files().get_media(fileId=file_id)
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)

            def download_in_thread():
                done = False
                while not done:
                    status, done = downloader.next_chunk()
                    if status:
                        logger.info(f"Прогресс загрузки: {int(status.progress() * 100)}%.")

            await asyncio.to_thread(download_in_thread)
            logger.success(f"Файл {file_id} успешно загружен.")
            file_io.seek(0)
            return file_io
        except Exception as e:
            logger.error(f"Критическая ошибка при попытке скачать файл из Google Drive: {e}", exc_info=True)
            raise e

    def _read_file_content(self, file: io.BytesIO, filename: str) -> str:
        """
        Читает содержимое файла (PDF, DOCX, TXT) и возвращает текст.
        """
        logger.info(f"Извлечение текста из файла: {filename}")
        try:
            if filename.lower().endswith('.pdf'):
                reader = PdfReader(file)
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            elif filename.lower().endswith('.docx'):
                doc = docx.Document(file)
                full_text = [para.text for para in doc.paragraphs if para.text.strip()]
                return "\n".join(full_text)
            else:
                return file.read().decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Ошибка при чтении файла {filename}: {e}")
            raise ValueError(f"Could not process file: {filename}")

    async def analyze_preparation(
            self,
            cv_file: io.BytesIO,
            cv_filename: str,
            feedback_text: str,
            requirements_link: str
    ) -> PreparationAnalysis:
        logger.info("Начало процесса оценки кандидата (Пайплайн 1)...")

        self._set_google_api_key()

        cv_text = self._read_file_content(cv_file, cv_filename)

        requirements_file_id = self._get_google_drive_file_id(requirements_link)
        requirements_text = await self._download_sheet_from_drive(requirements_file_id)

        session_service = InMemorySessionService()
        session_id = f"prep_session_{os.urandom(8).hex()}"
        user_id = "prep_user"
        await session_service.create_session(app_name=settings.app_name, user_id=user_id, session_id=session_id)

        logger.info("🚀 Запуск agent_1_data_parser...")
        runner_1 = Runner(agent=agent_1_data_parser, app_name=settings.app_name, session_service=session_service)
        message_for_agent_1 = types.Content(
            role="user",
            parts=[
                types.Part(text=f"cv_text: {cv_text}"),
                types.Part(text=f"requirements_text: {requirements_text}"),
                types.Part(text=f"feedback_text: {feedback_text}")
            ]
        )
        agent_1_output = ""
        async for event in runner_1.run_async(session_id=session_id, user_id=user_id, new_message=message_for_agent_1):
            if event.content and event.content.parts:
                agent_1_output += "".join(part.text for part in event.content.parts if part.text)

        logger.info("🚀 Запуск agent_2_grader...")
        runner_2 = Runner(agent=agent_2_grader, app_name=settings.app_name, session_service=session_service)
        message_for_agent_2 = types.Content(role="user", parts=[types.Part(text=agent_1_output)])
        agent_2_output = ""
        async for event in runner_2.run_async(session_id=session_id, user_id=user_id, new_message=message_for_agent_2):
            if event.content and event.content.parts:
                agent_2_output += "".join(part.text for part in event.content.parts if part.text)

        logger.info("🚀 Запуск agent_3_report_generator...")
        runner_3 = Runner(agent=agent_3_report_generator, app_name=settings.app_name, session_service=session_service)
        message_for_agent_3 = types.Content(role="user", parts=[types.Part(text=agent_2_output)])
        final_output = ""
        async for event in runner_3.run_async(session_id=session_id, user_id=user_id, new_message=message_for_agent_3):
            if event.content and event.content.parts:
                final_output += "".join(part.text for part in event.content.parts if part.text)

        logger.info("Парсинг финального вывода...")
        try:
            clean_json_str = final_output.strip().replace("```json", "").replace("```", "").strip()
            data = json.loads(clean_json_str)

            final_response_data = {
                "message": "Interview preparation report created successfully.",
                **data
            }
            return PreparationAnalysis(**final_response_data)
        except json.JSONDecodeError as e:
            logger.error(f"Ошибка декодирования JSON от Агента 3: {e}\nПолученный текст: {final_output}")
            raise ValueError("AI-сервис вернул некорректный формат данных.")
        except Exception as e:
            logger.error(f"Ошибка валидации Pydantic или другая ошибка: {e}")
            raise ValueError(f"Ошибка при формировании итогового ответа: {e}")

    async def _transcribe_audio_assemblyai(self, audio_data: io.BytesIO) -> str:
        logger.info("Начало транскрипции аудио через AssemblyAI с автоопределением языка...")

        config = aai.TranscriptionConfig(language_detection=True)
        transcriber = aai.Transcriber(config=config)

        def sync_transcribe_task():
            logger.info("Запуск синхронной задачи транскрипции в отдельном потоке...")
            return transcriber.transcribe(audio_data)

        transcript = await asyncio.to_thread(sync_transcribe_task)

        if transcript.status == aai.TranscriptStatus.error:
            logger.error(f"Ошибка транскрипции AssemblyAI: {transcript.error}")
            raise ValueError(f"Transcription failed: {transcript.error}")

        logger.success("Транскрипция AssemblyAI успешно завершена.")

        if not transcript.text:
            logger.warning("Транскрипция вернула пустой текст.")

        return transcript.text or ""

    async def analyze_results(self, video_link: str, matrix_content: bytes) -> ResultsAnalysis:
        logger.info("🚀 Запуск Пайплайна 2: Анализ результатов интервью...")

        self._set_google_api_key()

        file_id = self._get_google_drive_file_id(video_link)
        audio_file_stream = await self._download_audio_from_drive(file_id)
        transcription_text = await self._transcribe_audio_assemblyai(audio_file_stream)
        if not transcription_text:
            raise ValueError("Транскрипция не вернула текст.")

        session_service = InMemorySessionService()
        session_id = f"results_session_{os.urandom(8).hex()}"
        user_id = "results_user"
        await session_service.create_session(app_name=settings.app_name, user_id=user_id, session_id=session_id)

        logger.info("🚀 Запуск agent_4_topic_extractor...")
        runner_4 = Runner(agent=agent_4_topic_extractor, app_name=settings.app_name, session_service=session_service)
        message_for_agent_5 = types.Content(role="user", parts=[types.Part(text=transcription_text)])
        agent_5_output = ""
        async for event in runner_4.run_async(session_id=session_id, user_id=user_id, new_message=message_for_agent_5):
            if event.content and event.content.parts:
                agent_5_output += "".join(part.text for part in event.content.parts if part.text)

        logger.info("🚀 Запуск agent_5_final_report_generator...")
        runner_5 = Runner(agent=agent_5_final_report_generator, app_name=settings.app_name,
                          session_service=session_service)
        combined_input_for_agent_6 = f"### Транскрипция интервью:\n{transcription_text}\n\n### Матрица компетенций:\n{matrix_content.decode('utf-8', errors='ignore')}"
        message_for_agent_6 = types.Content(role="user", parts=[types.Part(text=combined_input_for_agent_6)])
        agent_6_output = ""
        async for event in runner_5.run_async(session_id=session_id, user_id=user_id, new_message=message_for_agent_6):
            if event.content and event.content.parts:
                agent_6_output += "".join(part.text for part in event.content.parts if part.text)

        logger.info("Парсинг и объединение финальных результатов...")
        try:
            clean_json_str_5 = agent_5_output.strip().replace("```json", "").replace("```", "").strip()
            topics_data = json.loads(clean_json_str_5)

            clean_json_str_6 = agent_6_output.strip().replace("```json", "").replace("```", "").strip()
            report_data = json.loads(clean_json_str_6)

            return ResultsAnalysis(
                message="Interview analysis completed successfully",
                transcription=transcription_text,
                scores=ScoreBreakdown(**report_data.get("scores", {})),
                strengths=report_data.get("strengths", []),
                concerns=report_data.get("concerns", []),
                recommendation=report_data.get("recommendation", "N/A"),
                reasoning=report_data.get("reasoning", ""),
                topicsDiscussed=topics_data.get("topicsDiscussed", [])
            )
        except json.JSONDecodeError as e:
            logger.error(f"Ошибка декодирования JSON: {e}")
            logger.error(f"Проблемный JSON от Агента 5: {agent_5_output}")
            logger.error(f"Проблемный JSON от Агента 6: {agent_6_output}")
            raise ValueError("AI-сервис вернул некорректный формат данных.")

    def create_docx_report(self, results: ResultsAnalysis) -> io.BytesIO:
        """Создает отчет в формате DOCX из результатов анализа."""

        document = Document()

        heading = document.add_heading('Отчет по результатам интервью', level=1)
        run = heading.runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(18)

        document.add_heading('Финальная рекомендация', level=2)
        p = document.add_paragraph()

        run_verdict_label = p.add_run("Вердикт: ")
        run_verdict_label.bold = True
        run_verdict_label.font.name = 'Calibri'
        run_verdict_label.font.size = Pt(12)

        run_verdict_text = p.add_run(results.recommendation)
        run_verdict_text.font.name = 'Calibri'
        run_verdict_text.font.size = Pt(12)

        p_reasoning = document.add_paragraph(results.reasoning)
        for run in p_reasoning.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)

        document.add_heading('Оценка компетенций', level=2)
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Компетенция'
        hdr_cells[1].text = 'Оценка (%)'

        for key, value in results.scores.model_dump().items():
            row_cells = table.add_row().cells
            row_cells[0].text = key.capitalize()
            row_cells[1].text = str(value)

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)

        document.add_heading('Ключевые сильные стороны', level=2)
        for strength in results.strengths:
            p_strength = document.add_paragraph(strength, style='List Bullet')
            for run in p_strength.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(12)

        document.add_heading('Области для улучшения', level=2)
        for concern in results.concerns:
            p_concern = document.add_paragraph(concern, style='List Bullet')
            for run in p_concern.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(12)

        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        return file_stream
